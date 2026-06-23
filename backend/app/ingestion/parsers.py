"""Document parsers package."""

from pathlib import Path


def get_parser(file_path: str):
    """Factory function to get the appropriate parser for a file type."""
    ext = Path(file_path).suffix.lower()
    
    parsers = {
        ".pdf": PDFParser,
        ".docx": DOCXParser,
        ".txt": TXTParser,
        ".csv": CSVParser,
        ".xlsx": XLSXParser,
        ".xls": XLSXParser,
    }
    
    parser_class = parsers.get(ext)
    if parser_class is None:
        raise ValueError(f"Unsupported file type: {ext}. Supported: PDF, DOCX, TXT, CSV, XLSX")
    
    return parser_class()


class BaseParser:
    """Base parser interface."""
    
    def parse(self, file_path: str) -> str:
        raise NotImplementedError


class PDFParser(BaseParser):
    """Parse PDF files using pypdf."""
    
    def parse(self, file_path: str) -> str:
        from pypdf import PdfReader
        
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n\n".join(text_parts)


class DOCXParser(BaseParser):
    """Parse DOCX files using python-docx."""
    
    def parse(self, file_path: str) -> str:
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)


class TXTParser(BaseParser):
    """Parse plain text files."""
    
    def parse(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()


class CSVParser(BaseParser):
    """Parse CSV files into structured text."""
    
    def parse(self, file_path: str) -> str:
        import pandas as pd
        
        df = pd.read_csv(file_path)
        text_parts = [f"CSV Data with {len(df)} rows and columns: {', '.join(df.columns.tolist())}"]
        
        # Convert each row to a readable format
        for row_num, (_, row) in enumerate(df.iterrows(), 1):
            row_text = " | ".join(f"{col}: {val}" for col, val in row.items() if pd.notna(val))
            text_parts.append(f"Row {row_num}: {row_text}")
        
        return "\n".join(text_parts)


class XLSXParser(BaseParser):
    """Parse Excel (.xlsx/.xls) files using openpyxl."""

    def parse(self, file_path: str) -> str:
        from openpyxl import load_workbook

        wb = load_workbook(file_path, read_only=True, data_only=True)
        text_parts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text_parts.append(f"Sheet: {sheet_name}")

            for row in ws.iter_rows():
                row_vals = [str(cell.value) if cell.value is not None else "" for cell in row]
                line = " | ".join(v for v in row_vals if v)
                if line.strip():
                    text_parts.append(line)

        wb.close()
        return "\n".join(text_parts)
